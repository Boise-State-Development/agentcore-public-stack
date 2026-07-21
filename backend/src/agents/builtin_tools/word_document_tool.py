"""Word document tools (create / modify / list / read).

Each tool runs python-docx code inside AWS Bedrock Code Interpreter and uses
the existing user-files store (``apis.shared.files``) for persistence and
delivery — generated/modified ``.docx`` files land in
``S3_USER_FILES_BUCKET_NAME`` with a ``FileMetadata`` row (status READY) in
``DYNAMODB_USER_FILES_TABLE_NAME``, so they appear in the chat's Files panel
and are downloadable via the app-api ``/files/{id}/preview-url`` route.

Tools
-----
* ``create_word_document`` — build a new document from python-docx code.
* ``modify_word_document`` — edit an existing document with python-docx code.
* ``list_word_documents``  — list the .docx files available in this chat.
* ``read_word_document``   — extract an existing document's text content.

(A page-screenshot/preview tool is intentionally omitted: rasterizing a
.docx requires LibreOffice/poppler, which the Python-only Code Interpreter
sandbox does not provide.)

Design notes
------------
* The Code Interpreter + user-files storage plumbing is shared with the Excel
  toolset and lives in ``builtin_tools.office._storage``; this module keeps
  only the python-docx specifics (preamble, generate/modify/extract) and the
  four tool factories.
* Identity (``user_id`` / ``session_id``) is captured by closure via the
  ``make_*`` factories — the same pattern used by the artifacts and
  spreadsheet_analysis tools (the Strands runtime here does NOT populate
  ``ToolContext.invocation_state`` with identity). The tools are injected
  per-request through ``extra_tools`` (see ``_build_word_document_tools`` in
  ``apis/inference_api/chat/routes.py``); they are deliberately NOT registered
  in ``builtin_tools/__init__`` because they need request-scoped identity.
"""

from __future__ import annotations

import asyncio
import logging
from typing import Any, Dict, Optional

from strands import tool

from agents.builtin_tools.office._storage import (
    _DocGenError,
    _ci_exec,
    _ci_read_bytes,
    _ci_write_bytes,
    _download_card,
    _download_s3_bytes,
    _error,
    _get_code_interpreter_id,
    _NO_CI_MESSAGE,
    _region,
    _storage_configured,
    _store_document,
    _validate_document_name,
)

logger = logging.getLogger(__name__)

# Word document MIME type (matches apis.shared.files.ALLOWED_MIME_TYPES).
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Sandbox path used to stage a source document loaded from S3.
_SANDBOX_SOURCE = "_source.docx"

_NO_STORAGE_MESSAGE = (
    "❌ Word document storage is not configured "
    "(S3_USER_FILES_BUCKET_NAME is not set on the runtime)."
)


# ---------------------------------------------------------------------------
# python-docx document builders (run in Code Interpreter)
# ---------------------------------------------------------------------------


_DOCX_PREAMBLE = (
    "from docx import Document\n"
    "from docx.shared import Pt, RGBColor, Inches\n"
    "from docx.enum.text import WD_ALIGN_PARAGRAPH\n"
)


def _generate_docx_bytes(
    code_interpreter_id: str, python_code: str, filename: str
) -> bytes:
    """Build a new .docx from user code and return its bytes.

    Blocking (boto3 / Code Interpreter) — call via ``asyncio.to_thread``.
    """
    from bedrock_agentcore.tools.code_interpreter_client import CodeInterpreter

    code_interpreter = CodeInterpreter(_region())
    code_interpreter.start(identifier=code_interpreter_id)
    try:
        # The user's code operates on a pre-initialized ``doc`` and must not
        # call Document()/doc.save() itself — we own the lifecycle.
        _ci_exec(
            code_interpreter,
            (
                f"{_DOCX_PREAMBLE}\n"
                "doc = Document()\n\n"
                f"{python_code}\n\n"
                f"doc.save({filename!r})\n"
            ),
        )
        data = _ci_read_bytes(code_interpreter, filename)
        if data is None:
            raise _DocGenError(
                f"Document '{filename}' was not produced. Make sure your code "
                "adds content to `doc`."
            )
        return data
    finally:
        try:
            code_interpreter.stop()
        except Exception:  # pragma: no cover - cleanup best-effort
            pass


def _modify_docx_bytes(
    code_interpreter_id: str,
    source_bytes: bytes,
    python_code: str,
    output_filename: str,
) -> bytes:
    """Load an existing .docx, apply user edits, return the new bytes.

    Blocking — call via ``asyncio.to_thread``.
    """
    from bedrock_agentcore.tools.code_interpreter_client import CodeInterpreter

    code_interpreter = CodeInterpreter(_region())
    code_interpreter.start(identifier=code_interpreter_id)
    try:
        _ci_write_bytes(code_interpreter, _SANDBOX_SOURCE, source_bytes)
        _ci_exec(
            code_interpreter,
            (
                f"{_DOCX_PREAMBLE}\n"
                f"doc = Document({_SANDBOX_SOURCE!r})\n\n"
                f"{python_code}\n\n"
                f"doc.save({output_filename!r})\n"
            ),
        )
        data = _ci_read_bytes(code_interpreter, output_filename)
        if data is None:
            raise _DocGenError(
                f"Modified document '{output_filename}' was not produced."
            )
        return data
    finally:
        try:
            code_interpreter.stop()
        except Exception:  # pragma: no cover - cleanup best-effort
            pass


def _extract_docx_text(code_interpreter_id: str, source_bytes: bytes) -> str:
    """Extract readable text (headings, paragraphs, tables) from a .docx.

    Blocking — call via ``asyncio.to_thread``.
    """
    from bedrock_agentcore.tools.code_interpreter_client import CodeInterpreter

    code_interpreter = CodeInterpreter(_region())
    code_interpreter.start(identifier=code_interpreter_id)
    try:
        _ci_write_bytes(code_interpreter, _SANDBOX_SOURCE, source_bytes)
        extraction = (
            "from docx import Document\n"
            f"doc = Document({_SANDBOX_SOURCE!r})\n"
            "lines = []\n"
            "for p in doc.paragraphs:\n"
            "    t = p.text.strip()\n"
            "    if not t:\n"
            "        continue\n"
            "    style = (p.style.name if p.style else '') or ''\n"
            "    if style.startswith('Heading'):\n"
            "        level = ''.join(ch for ch in style if ch.isdigit()) or '1'\n"
            "        lines.append('#' * min(int(level), 6) + ' ' + t)\n"
            "    else:\n"
            "        lines.append(t)\n"
            "for i, table in enumerate(doc.tables):\n"
            "    lines.append('')\n"
            "    lines.append('[Table %d]' % (i + 1))\n"
            "    for row in table.rows:\n"
            "        lines.append(' | '.join(c.text.strip() for c in row.cells))\n"
            "print('\\n'.join(lines))\n"
        )
        return _ci_exec(code_interpreter, extraction).strip()
    finally:
        try:
            code_interpreter.stop()
        except Exception:  # pragma: no cover - cleanup best-effort
            pass


# ---------------------------------------------------------------------------
# User-files lookup
# ---------------------------------------------------------------------------


async def _find_word_document(
    user_id: str, session_id: str, document_name: str
):
    """Find the newest READY .docx in this session matching ``document_name``.

    Returns the ``FileMetadata`` or ``None``. ``list_session_files`` returns
    newest-first, so the first match is the latest version.
    """
    from apis.shared.files import FileStatus, get_file_upload_repository

    target = (
        document_name
        if document_name.lower().endswith(".docx")
        else f"{document_name}.docx"
    )
    files = await get_file_upload_repository().list_session_files(
        session_id, status=FileStatus.READY
    )
    for meta in files:
        if (
            meta.user_id == user_id
            and meta.mime_type == _DOCX_MIME
            and meta.filename.lower() == target.lower()
        ):
            return meta
    return None


# ---------------------------------------------------------------------------
# Tool factories
# ---------------------------------------------------------------------------


def make_create_word_document_tool(session_id: str, user_id: str):
    """Create a ``create_word_document`` tool bound to the given identity."""

    @tool
    async def create_word_document(
        python_code: str,
        document_name: str,
    ) -> Any:
        """Create a new Word (.docx) document using python-docx code.

        Executes python-docx code in a sandboxed Code Interpreter to build a
        document from scratch, saves it to the user's files, and returns a
        download card. Great for structured reports with headings,
        paragraphs, tables, and matplotlib charts.

        Available libraries in the sandbox: python-docx, matplotlib, pandas,
        numpy.

        Args:
            python_code: python-docx code that builds the document. A blank
                document is already available as ``doc = Document()`` — do NOT
                call ``Document()`` or ``doc.save()`` yourself; the tool saves
                it for you. ``Pt``, ``RGBColor``, ``Inches`` and
                ``WD_ALIGN_PARAGRAPH`` are already imported.

                Example:
                    doc.add_heading('Quarterly Report', level=1)
                    doc.add_paragraph('Revenue increased by 15%...')
                    table = doc.add_table(rows=2, cols=2)
                    table.style = 'Light Grid Accent 1'
                    table.rows[0].cells[0].text = 'Quarter'

                To embed a chart, save a PNG with matplotlib then insert it:
                    import matplotlib.pyplot as plt
                    plt.figure(figsize=(8, 5))
                    plt.bar(['Q1', 'Q2'], [100, 120])
                    plt.savefig('chart.png', dpi=200, bbox_inches='tight')
                    plt.close()
                    doc.add_picture('chart.png', width=Inches(6))

            document_name: File name WITHOUT extension (.docx is added
                automatically). Use only letters, numbers, hyphens, and
                underscores (e.g. "sales-report", "Q4_analysis").

        Returns:
            An inline download card. The document is also saved to this
            chat's Files.
        """
        is_valid, error_msg = _validate_document_name(document_name)
        if not is_valid:
            return _error(
                f"❌ Invalid document name '{document_name}': {error_msg}\n\n"
                "Examples: sales-report, Q4_analysis, report-final"
            )

        filename = f"{document_name}.docx"
        code_interpreter_id = _get_code_interpreter_id()
        if not code_interpreter_id:
            return _error(_NO_CI_MESSAGE)
        if not _storage_configured():
            return _error(_NO_STORAGE_MESSAGE)

        try:
            file_bytes = await asyncio.to_thread(
                _generate_docx_bytes, code_interpreter_id, python_code, filename
            )
        except _DocGenError as exc:
            return _error(
                f"❌ Failed to create '{filename}'.\n\n```\n{exc}\n```\n\n"
                "Check the python-docx code for errors."
            )
        except Exception as exc:  # noqa: BLE001 - surface any sandbox error
            logger.error(f"create_word_document sandbox error: {exc}")
            return _error(f"❌ Failed to create '{filename}': {exc}")

        try:
            _id, download_url, size_kb = await _store_document(
                user_id, session_id, filename, file_bytes, _DOCX_MIME
            )
        except Exception as exc:  # noqa: BLE001 - storage failure is terminal
            logger.error(f"create_word_document storage error: {exc}")
            return _error(f"❌ Created '{filename}' but failed to save it: {exc}")

        return _download_card(filename, download_url, size_kb, "Created")

    return create_word_document


def make_modify_word_document_tool(session_id: str, user_id: str):
    """Create a ``modify_word_document`` tool bound to the given identity."""

    @tool
    async def modify_word_document(
        document_name: str,
        python_code: str,
        output_name: Optional[str] = None,
    ) -> Any:
        """Modify an existing Word (.docx) document with python-docx code.

        Loads a document previously created in this chat, runs your
        python-docx code against it, and saves the result (as a new file so
        the original is preserved). Returns a download card.

        Use ``list_word_documents`` first if you are unsure of the exact name.

        Args:
            document_name: Name of the existing document to edit (with or
                without the .docx extension), e.g. "sales-report".
            python_code: python-docx code that edits the document. The loaded
                document is available as ``doc = Document(...)`` — do NOT call
                ``Document()`` or ``doc.save()`` yourself. ``Pt``, ``RGBColor``,
                ``Inches`` and ``WD_ALIGN_PARAGRAPH`` are already imported.

                Example (append a section):
                    doc.add_heading('Addendum', level=1)
                    doc.add_paragraph('Updated figures for Q2.')

            output_name: Optional name (without extension) for the edited copy.
                Defaults to the source name (a new versioned copy is saved).

        Returns:
            An inline download card for the edited document.
        """
        code_interpreter_id = _get_code_interpreter_id()
        if not code_interpreter_id:
            return _error(_NO_CI_MESSAGE)
        if not _storage_configured():
            return _error(_NO_STORAGE_MESSAGE)

        source = await _find_word_document(user_id, session_id, document_name)
        if source is None:
            return _error(
                f"❌ No Word document named '{document_name}' was found in this "
                "chat. Use list_word_documents to see what's available."
            )

        out_base = output_name or source.filename
        if out_base.lower().endswith(".docx"):
            out_base = out_base[: -len(".docx")]
        is_valid, error_msg = _validate_document_name(out_base)
        if not is_valid:
            return _error(
                f"❌ Invalid output name '{out_base}': {error_msg}"
            )
        output_filename = f"{out_base}.docx"

        try:
            source_bytes = await asyncio.to_thread(
                _download_s3_bytes, source.s3_bucket, source.s3_key
            )
            file_bytes = await asyncio.to_thread(
                _modify_docx_bytes,
                code_interpreter_id,
                source_bytes,
                python_code,
                output_filename,
            )
        except _DocGenError as exc:
            return _error(
                f"❌ Failed to modify '{source.filename}'.\n\n```\n{exc}\n```\n\n"
                "Check the python-docx code for errors."
            )
        except Exception as exc:  # noqa: BLE001 - surface any sandbox error
            logger.error(f"modify_word_document error: {exc}")
            return _error(f"❌ Failed to modify '{source.filename}': {exc}")

        try:
            _id, download_url, size_kb = await _store_document(
                user_id, session_id, output_filename, file_bytes, _DOCX_MIME
            )
        except Exception as exc:  # noqa: BLE001 - storage failure is terminal
            logger.error(f"modify_word_document storage error: {exc}")
            return _error(
                f"❌ Modified '{source.filename}' but failed to save it: {exc}"
            )

        return _download_card(output_filename, download_url, size_kb, "Updated")

    return modify_word_document


def make_list_word_documents_tool(session_id: str, user_id: str):
    """Create a ``list_word_documents`` tool bound to the given identity."""

    @tool
    async def list_word_documents() -> Dict[str, Any]:
        """List the Word (.docx) documents available in this chat.

        Returns the file names and sizes of documents created or modified in
        this conversation. Use the names with modify_word_document or
        read_word_document.
        """
        from apis.shared.files import FileStatus, get_file_upload_repository

        files = await get_file_upload_repository().list_session_files(
            session_id, status=FileStatus.READY
        )
        seen: set[str] = set()
        rows = []
        for meta in files:  # newest-first
            if meta.user_id != user_id or meta.mime_type != _DOCX_MIME:
                continue
            if meta.filename in seen:
                continue
            seen.add(meta.filename)
            rows.append(f"- {meta.filename} ({meta.size_bytes / 1024:.1f} KB)")

        if not rows:
            text = (
                "No Word documents in this chat yet. Use create_word_document "
                "to make one."
            )
        else:
            text = "Word documents in this chat:\n" + "\n".join(rows)
        return {"content": [{"text": text}], "status": "success"}

    return list_word_documents


def make_read_word_document_tool(session_id: str, user_id: str):
    """Create a ``read_word_document`` tool bound to the given identity."""

    @tool
    async def read_word_document(document_name: str) -> Dict[str, Any]:
        """Read the text content of an existing Word (.docx) document.

        Extracts headings, paragraphs, and tables from a document created in
        this chat so you can reference or summarize its contents. Use
        list_word_documents first if unsure of the exact name.

        Args:
            document_name: Name of the document to read (with or without the
                .docx extension), e.g. "sales-report".

        Returns:
            The document's text content.
        """
        code_interpreter_id = _get_code_interpreter_id()
        if not code_interpreter_id:
            return _error(_NO_CI_MESSAGE)
        if not _storage_configured():
            return _error(_NO_STORAGE_MESSAGE)

        source = await _find_word_document(user_id, session_id, document_name)
        if source is None:
            return _error(
                f"❌ No Word document named '{document_name}' was found in this "
                "chat. Use list_word_documents to see what's available."
            )

        try:
            source_bytes = await asyncio.to_thread(
                _download_s3_bytes, source.s3_bucket, source.s3_key
            )
            text = await asyncio.to_thread(
                _extract_docx_text, code_interpreter_id, source_bytes
            )
        except _DocGenError as exc:
            return _error(f"❌ Failed to read '{source.filename}': {exc}")
        except Exception as exc:  # noqa: BLE001 - surface any sandbox error
            logger.error(f"read_word_document error: {exc}")
            return _error(f"❌ Failed to read '{source.filename}': {exc}")

        body = text or "(The document has no extractable text.)"
        return {
            "content": [
                {"text": f"Content of {source.filename}:\n\n{body}"}
            ],
            "status": "success",
        }

    return read_word_document
